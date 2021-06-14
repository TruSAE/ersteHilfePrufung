#!usr/bin/env python
# _*_ encoding: utf-8

from kivy.config import Config
from kivy.app import App
from kivy.lang import Builder
from kivy.properties import ObjectProperty
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooser
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.spinner import Spinner
from kivy.uix.textinput import TextInput

import json
from datetime import datetime
import docx
import os
import csv
import pandas as pd

Config.set('kivy', 'keyboard_mode', 'systemanddock')  # клавиатура мобильного приложения
Config.set('graphics', 'fullscreen', '1')


class StartScreen(Screen):
    """
    класс устанавливает внешний вид и методы справки и административного
    окна, закрывает работу программы
    """
    Builder.load_file('StartScreen.kv')

    def build_admin(self):
        Notebook.config_reader(self=Notebook)
        AdministerHuman.text_input_prufer = Notebook.member_prufer
        AdministerHuman.text_input_ort = Notebook.member_ort
        MainApp.sm.add_widget(AdministerHuman(name='administer_human'))
        MainApp.sm.current = 'administer_human'
        MainApp.sm.remove_widget(self)

    def build_note(self):
        MainApp.sm.add_widget(MyLayout(name='spinner'))

    def app_stop(self):
        App.get_running_app().stop()


class ChooseThemes(Screen):
    """
    класс устанавливает флаги выбора темы неполного экзамена
    """
    Builder.load_file('ChooseThemes.kv')

    # TODO обсчитать заново все суммы баллов по темам !
    def on_switch_first(self):
        if CheckExamThemes.first_visio_flag:
            CheckExamThemes.first_visio_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 19  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.first_visio_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 19  # исправление ошибочного свича

    def on_switch_recreation(self):
        if CheckExamThemes.recreation_flag:
            CheckExamThemes.recreation_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 7
        else:
            CheckExamThemes.recreation_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 7  # исправление ошибочного свича

    def on_switch_venosus(self):
        if CheckExamThemes.venosus_flag:
            CheckExamThemes.venosus_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 6  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.venosus_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 6  # исправление ошибочного свича

    def on_switch_arterial(self):
        if CheckExamThemes.arterial_flag:
            CheckExamThemes.arterial_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 10  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.arterial_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 10  # исправление ошибочного свича

    def on_switch_cpr(self):
        if CheckExamThemes.cpr_flag:
            CheckExamThemes.cpr_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 28  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.cpr_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 28  # исправление ошибочного свича

    def on_switch_fraktura_zakr(self):
        if CheckExamThemes.frakt_zakr_flag:
            CheckExamThemes.frakt_zakr_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 10  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.frakt_zakr_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 10  # исправление ошибочного свича

    def on_switch_fraktura_otkr(self):
        if CheckExamThemes.frakt_otkr_flag:
            CheckExamThemes.frakt_otkr_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 15  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.frakt_otkr_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 15  # исправление ошибочного свича

    def on_switch_heiml_orto(self):
        if CheckExamThemes.heiml_otro_flag:
            CheckExamThemes.heiml_otro_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 14  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.heiml_otro_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 14  # исправление ошибочного свича

    def on_switch_heiml_down(self):
        if CheckExamThemes.heiml_down_flag:
            CheckExamThemes.heiml_down_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 11  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.heiml_down_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 11  # исправление ошибочного свича

    def on_switch_transport(self):
        if CheckExamThemes.transport_flag:
            CheckExamThemes.transport_flag = False
            CheckExamThemes.result_flag += 1
            Notebook.part_MAXIMUM -= 7  # уменьшение общей макисмальной суммы штрафов
        else:
            CheckExamThemes.transport_flag = True
            CheckExamThemes.result_flag -= 1
            Notebook.part_MAXIMUM += 7  # исправление ошибочного свича

    def choose_yes(self):
        if Notebook.part_MAXIMUM != 0:
            MainApp.sm.add_widget(CheckExamThemes(name='check_exam_themes'))
            MainApp.sm.current = 'check_exam_themes'
            MainApp.sm.remove_widget(self)
            AdministerHuman.name_flag = False  # опущен флаг имени.
            AdministerHuman.number_flag = False  # опущен флаг номера группы
            Notebook.summ_result = 0  # обнуляем суммарный результат

    def choose_no(self):  #
        MainApp.sm.remove_widget(self)
        MainApp.sm.current = 'administer_human'
        MainApp.sm.remove_widget(self)
        AdministerHuman.choose_themes_widget(self=AdministerHuman)


class AdministerHuman(Screen):
    """
    класс содержит методы записи параметров экзаменатора, студента, группы, устаналивает
    переход к полному (стандартному) или частичному экзаменам, начинает экзамен,
    закрывает работу программы
    """
    Builder.load_file('AdministerHuman.kv')

    name_student = ''
    text_input_name = ObjectProperty()  # окно ввода фамилии студента
    name_flag = False  # флаг записи студента
    text_number_grupp = ObjectProperty()  # окно ввода номера группы
    number_flag = False  # флаг записи номера
    text_input_prufer = ''  # имя экзаменатора
    text_input_ort = ''  # место проведния экзамена

    def write_ort_and_prufer(self):
        """Проверка введения имени экзаменатора и места проведения. Запись экзамена и места словарь"""
        if self.display_ort.text.strip() != '' and self.display_prufer.text.strip() != '':
            if self.display_ort.text in Notebook.ort:
                self.text_input_ort = self.display_ort.text
                Notebook.data['ort'] = self.text_input_ort.strip()
            else:
                Notebook.ort.append(self.display_ort.text)
                self.text_input_ort = self.display_ort.text
                Notebook.data['ort'] = self.text_input_ort
            if self.display_prufer.text in Notebook.prufer:
                self.text_input_prufer = self.display_prufer.text
                Notebook.data['prufer'] = self.text_input_prufer.strip()
            else:
                Notebook.prufer.append(self.display_prufer.text)
                self.text_input_prufer = self.display_prufer.text
                Notebook.data['prufer'] = self.text_input_prufer

    def clear_ort_and_instruktor(self):
        """Очистка полей ввода экзаменатора и места проведения"""
        self.display_ort.text = ''
        self.display_prufer.text = ''

    def write_number(self):
        """Проверка введения номера группы. Запись номера группы в словарь"""
        if self.text_number_grupp.text != '':
            Notebook.number_grupp = self.text_number_grupp.text  # записываем номер группы в класс записная книжка
            Notebook.data['number_grupp'] = self.text_number_grupp.text  # записываем номер группы в словарь
            self.number_flag = True  # флаг записи номера поднят. номер в базе

    def write_name(self):
        """Проверка введения фамилии студента. Запись фамилии в словарь"""
        if self.text_input_name.text != '':  # если введено новое имя студента
            Notebook.name_student = self.text_input_name.text  # запись студента в ноутбук
            Notebook.data[
                'name'] = self.text_input_name.text.title()  # сохранение фамилии в словарь. прописные буквы в словах

            self.name_flag = True  # поднят флаг студента

    def choose_themes_widget(self):  # переключение между частичным и полным экзаменами
        """устанаввливает флаги каждой темы в True после проведённого частичного экзамена
        и опускает флаг выбора частичного экзамена"""
        CheckExamThemes.transport_flag = True
        CheckExamThemes.cpr_flag = True
        CheckExamThemes.arterial_flag = True
        CheckExamThemes.heiml_down_flag = True
        CheckExamThemes.heiml_otro_flag = True
        CheckExamThemes.frakt_otkr_flag = True
        CheckExamThemes.frakt_zakr_flag = True
        CheckExamThemes.venosus_flag = True
        CheckExamThemes.recreation_flag = True
        CheckExamThemes.first_visio_flag = True
        CheckExamThemes.result_flag = 0

    def app_stop(self):  # TODO потом убрать ? везде убрать ?
        App.get_running_app().stop()

    def choose_spin_builder(self):
        """Вызов спиннера выбора экзаменатора и места проведения при нескольких различных таких"""
        box_content = BoxLayout(orientation='vertical')
        spinner = Spinner(text='Место проведения экзамена', values=[value for value in Notebook.ort])
        spinner_1 = Spinner(text='Экзаменатор', values=[value for value in Notebook.prufer])
        btn_yes = Button(text='Подтвердить выбор', on_press=(lambda *args: yes()))
        box_content.add_widget(spinner)
        box_content.add_widget(spinner_1)
        box_content.add_widget(btn_yes)
        box_popup = Popup(title='Место действия и экзаменатор',
                          content=box_content,
                          size_hint=(0.8, 0.4))
        box_popup.open()

        def yes():
            self.display_ort.text = spinner.text
            Notebook.member_ort = spinner.text
            Notebook.data['ort'] = Notebook.member_ort
            self.display_prufer.text = spinner_1.text
            Notebook.member_prufer = spinner_1.text
            Notebook.data['prufer'] = Notebook.member_prufer
            if spinner.text != 'Место проведения экзамена' and spinner_1.text != 'Экзаменатор':
                box_popup.dismiss()

    def write_parametres(self):
        """Запись в словарь экзаменатора, места проведения, уникального номера студента.
        Установка параметров имени, группы студента для каждого окна"""
        CheckExamThemes.name_student = Notebook.name_student
        CheckExamThemes.number_grupp = Notebook.number_grupp
        Notebook.data['prufer'] = self.text_input_prufer
        Notebook.data['ort'] = self.text_input_ort
        Notebook.id_student_plus(self=Notebook)
        Notebook.data['ID'] = Notebook.id_student
        Notebook.config_writer(self=Notebook)

    def build(self):
        if self.number_flag and self.name_flag:  # проверка введения места и экзаменатора
            if self.ids.partial_switch.active:  # если выбран частичный экзамен...
                Notebook.choose_themes_flag = True
                Notebook.part_MAXIMUM = 127  # возвращаем частичный максимум к исходному значению
                self.write_parametres()
                MainApp.sm.add_widget(ChooseThemes(name='choose_themes'))  # переход к выбору отдельных тем
                MainApp.sm.current = 'choose_themes'
                self.text_input_name.text = ''  # для очистки поля ввода имени
                self.ids.partial_switch.active = False  # перевод свича в исходное состояние

            else:
                self.write_parametres()
                MainApp.sm.add_widget(CheckExamThemes(name='check_exam_themes'))  # к полному экзамену
                MainApp.sm.current = 'check_exam_themes'
                self.text_input_name.text = ''  # для очистки поля ввода имени
                self.name_flag = False  # флаг имени опущен
                self.number_flag = False  # флаг номера группы опущен
                Notebook.choose_themes_flag = False  # опускаем флаг выбора частичного экзамена
                Notebook.summ_result = 0  # обнуляем суммарный результат
        else:
            self.vvarning_popup()

    def vvarning_popup(self):
        box_content = BoxLayout(orientation='vertical')
        label_warning = Label(text='Запишите сведения о студенте.')
        btn_yes = Button(text='Закрыть окно', on_press=(lambda *args: box_popup.dismiss()))
        box_content.add_widget(label_warning)
        box_content.add_widget(btn_yes)
        box_popup = Popup(title='Предупреждение!', content=box_content, size_hint=(0.9, 0.2))
        box_popup.open()

        #def dismiss():
         #   box_popup.dismiss()

    def build_analis(self):
        Analis.real_date(self=Analis)  # установка текущей даты
        Analis.grafics(self=Analis)  # предварительная обработка данных и подготовка фрейма к анализу
        MainApp.sm.add_widget(Analis(name='analis'))
        MainApp.sm.current = 'analis'


class MyFileChooser(Screen):#, FileChooser):
    Builder.load_file('MyFileChooser.kv')

    def path_label_text(self, *args):
        label_text = 'Сохранить итоговый документ в: '
        self.ids.path_label.text = label_text + ''.join(str(args[1])).strip("['").strip("']")

    def save(self, path):
        Notebook.path_saving_result_and_recomendation = path
        if Notebook.grupp_analis_flag:
            Analis.print_result_grupp(self=Analis)
        elif Notebook.name_fehler_flag:
            Analis.print_fehler_name(self=Analis)
        elif Notebook.date_analys_flag:
            Analis.print_result_date(self=Analis)

    def cancel_saving(self):
        MainApp.sm.remove_widget(self)
        MainApp.sm.current = 'analis'


class Analis(Screen):
    """
    Класс содержит методы графического представления результатов прошедших сессий,
    формирования отчётного файла результатов экзамена, его сохранение, печать
    или отправку по адресу электронной почты
    """
    #  TODO навесить календарь из питона с выводом ...на календарь дат из словаря. вывод в label статистики.
    Builder.load_file('Analis.kv')

    mean_res = ''  # хранение результатов поиска и анализа, для отображения на метке
    calendar_date_input = ''  # отображение текущей даты в окне ввода даты
    date_mask = ''  # дата для поиска в датафрейме
    label_res = ''  # строка на label принимает mean_res
    name_input = ''  # строка для ввода имени для поиска по имени студента
    name_mask = ''  # имя для поиска в датафрейме

    date_mask_popup = ''

    def file_chooser(self):
        """метод порождает новое окно выбора места сохранения итогового документа"""
        MainApp.sm.add_widget(MyFileChooser(name='my_file_chooser'))
        MainApp.sm.current = "my_file_chooser"

    def analysis_resultatis_fehler(self, res_grupp):  # TODO принимает уже готовый фрейм!!!!
        # TODO теряет ошибки!!!!! пока исключить этот метод из релиза до отработки
        """
        анализ предсуществующего фрейма (выборка по дате, имени или группе) на
        содержание ошибок в каждой теме экзамена, если  теме были ошибки
        :param res_grupp: pandas.DataFrame, созданный в методах поиска по дате, имени или группе
        :return: файл типа json TODO сделать выбор вывода на экран или сохранения на диск
        """
        fehler_res_grupp = res_grupp[res_grupp.summ_fehler != 0]
        res_fehler_items = {}
        keys_fehler = {}
        items_fehler = []  # для названий ошибки
        keyss = []  # список для названий темы ошибки
        for i in range(fehler_res_grupp.shape[0]):  # итерация по всем строкам
            for j in range(7, fehler_res_grupp.shape[1]):  # итерация по столбцам учёта ошибок
                # если есть баллы ошибок и это не корригирующие баллы
                if fehler_res_grupp.iloc[i, j] != 0 and \
                        fehler_res_grupp.columns[j].split('.')[0].strip() != 'Корректирующие баллы':
                    if len(keyss) == 0:  # если список названий тем экзамена пуст
                        # добавляем название темы
                        keyss.append(fehler_res_grupp.columns[j].split('.')[0].strip())
                        # добавляем название ошибки
                        items_fehler.append(fehler_res_grupp.columns[j].split('.')[1].strip())
                    elif keyss[0] == fehler_res_grupp.columns[j].split('.')[0].strip():  # если названия тем идентичны
                        # добавляем название ошибки
                        items_fehler.append(fehler_res_grupp.columns[j].split('.')[1].strip())
                    else:  # если название темы уже другое, т.е. прошли по всем ошибкам темы экзамена
                        # добавляем в словарь ключ тема : значение
                        keys_fehler[keyss[0]] = items_fehler.copy()
                        # список ошибок через копирование, иначе - пустой
                        keyss.clear()  # очищаем список тем экзамена
                        items_fehler.clear()  # очищаем список ошибок
                        continue  # вновь идём по кругу
            # вышли из цикла столбцов и добавляем в словарь ключ имя:
            res_fehler_items[fehler_res_grupp.iloc[i, 0]] = keys_fehler.copy()
            # значение словарь ошибок
            keys_fehler.clear()  # очищаем словарь ключей и ошибок
        print(json.dumps(res_fehler_items, ensure_ascii=False, sort_keys=True,
                         indent=4))  # имя, тема и ошибки. красивый вывод json

    def real_date(self):
        """
        установка текущей даты в окно ввода даты поиска
        """
        z = datetime.now()
        self.calendar_date_input = str(datetime.strftime(z, '%d.%m.%Y'))

    def find_name(self):
        """Поиск по фамилии, имени и отчеству, вывод на экран результатов экзамена: даты, оценки, штрафных баллов,
        значимых ошибок (более 3-х баллов), подготовка датафрейма для проведения анализа в методе печати"""
        if os.path.exists('students_data.csv'):
            self.name_mask = self.ids.name_input.text.strip().title()  # маска имени
            self.date_mask_popup = ""
            if self.name_mask not in Notebook.result['name'].values:  # если нет имени в фрейме
                self.mean_res = f'не обнаружено {self.name_mask}'.center(20)
                self.printing_res_label()
                # если имён более 1. Запрос даты и проверка формата ввода даты
            elif Notebook.result[Notebook.result['name'] == self.name_mask].shape[0] > 1:
                self.date_input_popup()  # всплывающее окно выбора даты и дальнейшего анализа
            else:  # выборка по имени когда имя одно
                Notebook.res_name = Notebook.result[Notebook.result['name'] == self.name_mask]
                self.mean_res = f'Для записи {self.name_mask.title()}:\n'.center(20)
                self.mean_res += f'Оценка на экзамене {Notebook.res_name.iloc[0, 3]}\n'.center(20)
                self.mean_res += f'Штрафных баллов {Notebook.res_name.iloc[0, 4]}\n'.center(20)
                if Notebook.res_name.iloc[0, 4] != 0:
                    Notebook.name_fehler_flag = True  # ошибки найдены
                    self.mean_res += 'Значимые ошибки:\n'.center(20)
                    self.fehler_analys()  # вызов метода обработки полученных ошибок для вывода на экран
                self.printing_res_label()
                self.ids.btn_saving.disabled = False

    def print_fehler_name(self):  # формирование документа с таблищей ошибок и сохранение на диск
        """
        Создание документа "*.docx", содержащего сведения о слушателе, экзаменаторе,
        совершенных ошибках и коррекционных ворпросах
        """
        if Notebook.name_fehler_flag:  # при обнаруженных ошибках в find_name()
            doc = docx.Document()
            # добавляем первый параграф (абзац)
            doc.add_paragraph(Notebook.res_name.iloc[0, 6])  # ort из find_name()
            # добавляем параграф
            doc.add_paragraph('Список зафиксированных ошибок и ответов на дополнительные вопросы при практическом '
                              'испытании по дисциплине')
            doc.add_paragraph('"Оказание первой помощи"')
            # добавляем параграф
            par1 = doc.add_paragraph('Группа № ')
            # добавляем текст в третий параграф
            par1.add_run(str(Notebook.res_name.iloc[0, 1])).bold = True  # number_grupp
            # добавляем текст в третий параграф
            par2 = doc.add_paragraph('Фамилия ')
            par2.add_run(str(Notebook.res_name.iloc[0, 0])).bold = True  # добавление имени
            # добавляем параграф
            par3 = doc.add_paragraph('Оценка на экзамене ')
            par3.add_run(str(Notebook.res_name.iloc[0, 3])).bold = True  # добавление оценки
            doc.add_paragraph('')
            dct_fehler_name = {}  # словарь для темы и значения ошибок
            keys_lst_fehler = []  # список для тем ошибок
            items_lst_fehler = []  # список для значения ошибок
            for i in range(8, 74):  # итерация по всем столбцам фрейма, где значение больше 0
                if Notebook.res_name.iloc[0, i] > 0:
                    if len(keys_lst_fehler) == 0:  # если список тем пуст
                        keys_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[0])  # добавляем тему
                        items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем значение
                    elif keys_lst_fehler[0] == Notebook.res_name.columns[i].split('. ')[0]:
                        items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем значение
                    else:  # если новое название темы
                        dct_fehler_name[keys_lst_fehler[0]] = items_lst_fehler.copy()  #
                        # записываем в словарь пару тема/список значений
                        keys_lst_fehler.clear()  # очищаем списки
                        items_lst_fehler.clear()
                        keys_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[0])  # добавляем новую тему
                        items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем новое
                        # значение
                    dct_fehler_name[keys_lst_fehler[0]] = items_lst_fehler.copy()  # записываем в
                    # словарь пару
                    # тема: список значений для последней итерации
            keys_lst_fehler.clear()  # очищаем списки
            items_lst_fehler.clear()
            for key in dct_fehler_name.keys():
                doc.add_paragraph(key + ':')
                doc.add_paragraph('; '.join(dct_fehler_name[key]))
                doc.add_paragraph('-' * 100)

            doc.add_paragraph('')
            # добавляем таблицу с именем экзаменатора
            table3 = doc.add_table(rows=1, cols=2)
            table3.cell(0, 0).text = "Экзаменатор"
            table3.cell(0, 1).text = Notebook.res_name.iloc[0, 5]
            doc.add_paragraph('')
            # добавляем таблицу с датой создания документа
            doc.add_paragraph(f"{datetime.strftime(datetime.now(), '%d.%m.%Y')}")  # печать текущей даты
            try:
                doc.save(
                    f'{Notebook.path_saving_result_and_recomendation}/mistakes_{str(Notebook.res_name.iloc[0, 0])}_{str(Notebook.res_name.iloc[0, 2])}.docx')  # вызывается в окне выбора папки сохранения

            except:
                self.permission_denied(self)
            Notebook.name_fehler_flag = False  # опустить флаг ошибок

    def date_input_popup(self):
        """всплывающее окно для предупреждения о нескольких совпадениях при поиске фамилии"""
        lst_date = list(set(Notebook.result[Notebook.result['name'] == self.name_mask].date.values))
        text_find_date = 'Найдены следующие даты:\n'
        date_from_df = ', '.join(lst_date).strip(', ')
        box_content = BoxLayout(orientation='vertical')
        label_date = Label(text=text_find_date + date_from_df)
        label_warning = Label(text='Запишите требуемую дату в формате дд.мм.гггг')
        date_input_popup = TextInput()
        btn_yes = Button(text='Подтвердить',
                         on_press=(lambda *args: [date_input_find_name(date_input_popup.text), box_popup.dismiss()]))
        btn_no = Button(text='Закрыть окно', on_press=(lambda *args: box_popup.dismiss()))
        box_content.add_widget(label_date)
        box_content.add_widget(label_warning)
        box_content.add_widget(date_input_popup)
        box_btns = BoxLayout(orientation='horizontal')
        box_btns.add_widget(btn_no)
        box_btns.add_widget(btn_yes)
        box_content.add_widget(box_btns)
        box_popup = Popup(title='Множественные совпадения в одной или нескольких датах!',
                          content=box_content, size_hint=(0.9, 0.4))
        box_popup.open()

        #def popup_dismiss():
         #   box_popup.dismiss()

        def date_input_find_name(*args):  # поиск одного имени среди нескольких дат
            """поиск по введённой дате при нескольких совпадениях и вывод на экран"""
            self.date_mask_popup = str(args[0])
            # res_name_multi_date новый фрейм с двумя условиями
            Notebook.res_name = Notebook.result[
                (Notebook.result.name == self.name_mask) & (Notebook.result.date == self.date_mask_popup)]
            if self.date_mask_popup != '' and self.date_mask_popup in Notebook.res_name['date'].values:
                self.mean_res = f'За {self.date_mask_popup} для {self.name_mask.title()}:\n'.center(20)
                self.mean_res += f'Оценка на экзамене {Notebook.res_name.iloc[0, 3]}\n'.center(20)
                self.mean_res += f'Штрафных баллов {Notebook.res_name.iloc[0, 4]}\n'.center(20)
                self.mean_res += 'Значимые ошибки:\n'.center(20)
                self.fehler_analys()  # создание словаря тем и ошибок для вывода на экран
                self.printing_res_label()
                self.ids.btn_saving.disabled = False  # разблокироать кнопку "сохранить"

            else:  # не встречается
                self.mean_res = f'Записи слушателя по фамилии:\n'.center(20)
                self.mean_res += f'{self.name_mask.title()} не обнаружены\n'.center(20)
                self.mean_res += f'среди записей за {self.date_mask_popup}'.center(20)
                self.printing_res_label()

    def fehler_analys(self):
        """анализ значимых ошибок (3 и более штрафных балла за ошибку), вывод на экран"""
        # обработка ошибок и вывод на экран
        dct_fehler_name = {}  # словарь для темы и значения ошибок
        keys_lst_fehler = []  # список для тем ошибок
        items_lst_fehler = []  # список для значения ошибок
        for i in range(8, 74):  # итерация по всем столбцам фрейма, где значение больше 3
            if Notebook.res_name.iloc[0, i] >= 3:
                if len(keys_lst_fehler) == 0:  # если список тем пуст
                    keys_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[0])  # добавляем тему
                    items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем значение
                elif keys_lst_fehler[0] == Notebook.res_name.columns[i].split('. ')[0]:
                    items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем значение
                else:  # если новое название темы
                    dct_fehler_name[keys_lst_fehler[0]] = items_lst_fehler.copy()  #
                    # записываем в словарь пару тема/список значений
                    keys_lst_fehler.clear()  # очищаем списки
                    items_lst_fehler.clear()
                    keys_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[0])  # добавляем новую тему
                    items_lst_fehler.append(Notebook.res_name.columns[i].split('. ')[1])  # добавляем новое
                    # значение
                dct_fehler_name[keys_lst_fehler[0]] = items_lst_fehler.copy()  # записываем в
                # словарь пару
                # тема: список значений для последней итерации
        keys_lst_fehler.clear()  # очищаем списки
        items_lst_fehler.clear()
        for key in dct_fehler_name:  # записываем в словарь пару тема/список значений
            self.mean_res += f'{key}:\n{" ".join(dct_fehler_name[key])}\n\n'.center(20)
        dct_fehler_name.clear()  # очистка словаря перед новой записью

    def find_grupp(self, inpt_grupp):
        """
        проверка правильности введённого номера группы и наличие этого номера во фрейме
        при отсутствии сведений или неправильном формате - предупреждения на экране
        :return: отображение статистики результатов поиска на экране
        """
        if os.path.exists('students_data.csv'):
            try:
                grupp_mask = int(inpt_grupp)  # маска номера группы
                if grupp_mask not in Notebook.result['number_grupp'].values and \
                        len(Notebook.result['number_grupp'].values) != 0:
                    lst_number_grupp = list(set(Notebook.result.number_grupp.values))

                    self.mean_res = f'За период с {Notebook.result.date.min()} по {Notebook.result.date.max()}\n'.center(20)
                    self.mean_res += f'существуют записи для {len(lst_number_grupp)} групп.\n'.center(20)
                    self.mean_res += f'Выберите номер группы из списка:\n'.center(20)
                    self.mean_res += f"{', '.join(str(i) for i in lst_number_grupp).strip(', ')}".center(20)
                    self.printing_res_label()
                else:
                    Notebook.res_grupp = Notebook.result[Notebook.result.number_grupp == grupp_mask]  # выборка по группе
                    ohne_fehler = Notebook.res_grupp[Notebook.res_grupp['summ_fehler'] == 0]  # выборка без ошибок в группе

                    ohne_bewertung = Notebook.res_grupp[(Notebook.res_grupp['оценка'] == 2) | (Notebook.res_grupp[
                                                                                                   'оценка'] == 0)].оценка.count()  # кол-во
                    # людей в группе без аттестации (получили 2 или без оценки)
                    bewertung_4 = Notebook.res_grupp[
                        Notebook.res_grupp['оценка'] == 4].оценка.count()  # кол-во людей в группе получили 4
                    bewertung_3 = Notebook.res_grupp[
                        Notebook.res_grupp['оценка'] == 3].оценка.count()  # кол-во людей в группе получили 3

                    bewertung_5 = Notebook.res_grupp.оценка[
                        Notebook.res_grupp['оценка'] == 5].count()  # получили 5. включить в отчёт неявно
                    fehler_res_grupp = Notebook.res_grupp[Notebook.res_grupp.summ_fehler != 0]  # выборка с ошибками
                    # сколько человек получили 5 с ошибками. включить
                    leute_5_mit_fehler = Notebook.res_grupp[(Notebook.res_grupp.оценка == 5) & (
                            Notebook.res_grupp.summ_fehler != 0)].name.count()
                    # в отчёт неявно
                    # анализ ошибок группы  TODO куда выводить? отключено
                    # self.analysis_resultatis_fehler(Notebook.res_grupp)
                    Notebook.grupp_analis_flag = True  # поднять флаг проведённого анализа
                    # формирование строки вывода итого на экран
                    self.mean_res = f'Для группы № {grupp_mask} найдено записей: {Notebook.res_grupp.name.count()}.\n'.center(
                        20)
                    self.mean_res += f'Средняя сумма ошибок: {round(Notebook.res_grupp.summ_fehler.mean(), 2)}\n'.center(20)
                    self.mean_res += f'Получили 5: {bewertung_5}\n'.center(20)
                    self.mean_res += f'Получили 4: {bewertung_4}\n'.center(20)
                    self.mean_res += f'Получили 3: {bewertung_3}\n'.center(20)
                    self.mean_res += f'Средняя оценка по группе: {round(Notebook.res_grupp.оценка.mean(), 2)}\n'.center(20)
                    self.mean_res += f'Итогов с ошибками: {fehler_res_grupp.summ_fehler.count()}\n'.center(20)
                    self.mean_res += f'Итогов без ошибок: {ohne_fehler.summ_fehler.count()}\n'.center(20)
                    self.mean_res += f'Получили 5 с ошибками: {leute_5_mit_fehler}\n'.center(20)
                    self.mean_res += f'Не аттестовано {ohne_bewertung}'.center(20)
                    self.printing_res_label()
                    self.ids.btn_saving.disabled = False
            except:
                self.mean_res = 'Подразумевается числовое значение номера группы'.center(20)
                self.printing_res_label()

    def find_date(self):
        """
        проверка правильности введённой даты поиска и наличие этой даты во фрейме
        при отсутствии сведений или неправильном формате даты - окно предупреждения
        :return: отображение статистики результатов поиска на экране
        """
        if os.path.exists('students_data.csv'):

            if len(self.display_date_input.text) != 10:
                self.false_date_popup()  # всплывающее предупреждение
            else:
                try:
                    datetime.strptime(self.display_date_input.text, '%d.%m.%Y')  # проверить формат даты
                    self.date_mask = self.display_date_input.text
                    if self.date_mask not in Notebook.result['date'].values:  # дата не найдена
                        self.no_date_popup()  # всплывающее предупреждение
                    else:
                        Notebook.res_date = Notebook.result[Notebook.result['date'] == self.date_mask]  # выборка на дату
                        Notebook.name_analis = Notebook.res_date.sort_values('name')  # сортировка по именам
                        Notebook.name_res = Notebook.name_analis[['name', 'оценка']]  # взять только столбцы
                        self.mean_res = f"{Notebook.res_date.name.count()} слушателей сдавал(и) экзамен {self.date_mask}.\n".center(
                            20)
                        self.mean_res += f"{Notebook.res_date[Notebook.res_date['summ_fehler'] == 0].name.count()} сдали экзамен без ошибок,\n".center(
                            20)
                        self.mean_res += f"{Notebook.res_date.оценка[Notebook.res_date['оценка'] == 5].count()} получили оценку 5,\n".center(
                            20)
                        self.mean_res += f"{Notebook.res_date[Notebook.res_date['оценка'] == 4].оценка.count()} получили оценку 4,\n".center(
                            20)
                        self.mean_res += f"{Notebook.res_date[Notebook.res_date['оценка'] == 3].оценка.count()} получили оценку 3,\n".center(
                            20)
                        self.mean_res += f"не аттестовано {Notebook.res_date[(Notebook.res_date['оценка'] == 2) | (Notebook.res_date['оценка'] == 0)].оценка.count()}.".center(
                            20)
                        self.printing_res_label()
                        Notebook.date_analys_flag = True  # флаг анализ по дате поднят для сохранения
                        self.ids.btn_saving.disabled = False  # актиаия кнопки "сохранять" результаты поиска по дате
                except:
                    self.false_date_popup()  # всплывающее предупреждение

    def no_date_popup(self):
        '''
        Всплывающее окно с кнопкой
        content: нет сведений на выбранную дату
        '''
        box_content = BoxLayout(orientation='vertical')
        label_warning = Label(text='Для выбранной даты нет сведений')
        btn_yes = Button(text='Закрыть окно', on_press=(lambda *args: box_popup.dismiss()))
        box_content.add_widget(label_warning)
        box_content.add_widget(btn_yes)
        box_popup = Popup(title='Ошибка поиска!', content=box_content, size_hint=(0.9, 0.4))
        box_popup.open()

    def grafics(self):
        """
        проверка наличия и предварительный анализ .csv-файла с сохранёнными данными
        срабатывает автоматически при входе в окно
        :return: Notebook.result = pandas.DataFrame
        """
        if os.path.exists('students_data.csv'):
            res = pd.read_csv('students_data.csv', index_col=False)
            res.index = res['ID']
            Notebook.result = res.drop('ID', axis=1)
        else:
            box_content = BoxLayout(orientation='vertical')
            label_warning = Label(text='Пока нет данных для анализа и отображения')
            box_content.add_widget(label_warning)
            btn_close = Button(text='Закрыть', size_hint=(1, 0.4), on_press=(lambda *args: box_popup.dismiss()))
            box_content.add_widget(btn_close)
            box_popup = Popup(title='Ошибка поиска!', content=box_content, size_hint=(0.8, 0.4))
            box_popup.open()


    def printing_res_label(self):
        """вывод в Label форматированного текста результатов поиска и анализа"""
        self.ids.label_res.text = self.mean_res

    def exit_analis(self):
        """Выход из окна анализа, уничтожение этого окна и переход в окно администрирования"""
        MainApp.sm.remove_widget(self)
        MainApp.sm.current = 'administer_human'

    def calendar(self):  # отключено
        pass

    def disabled_btn_saving(self):  # деактивация кнопки
        self.ids.btn_saving.disabled = True

    def print_result_grupp(self):
        """
        Создание документа "*.docx", содержащего сведения о месте экзамена, экзаменаторе
        фамилии, имена и отчества студентов, сдававших экзамен на определённую дату
        (по умолчанию, на текущую дату) с полученными оценками.
        """
        if Notebook.grupp_analis_flag:  # при созданном и непустом фрейме, проведённом анализе find_grupp()
            doc = docx.Document()
            # добавляем первый параграф (абзац)
            Notebook.name_analis = Notebook.res_grupp.sort_values(
                'name')  # сортировка по именам
            Notebook.name_res = Notebook.name_analis[['name', 'оценка']]  # взять только столбцы
            doc.add_paragraph(Notebook.name_analis.iloc[0, 6])  # ort из name_analis из find_date()
            # добавляем параграф
            par1 = doc.add_paragraph('Аттестационный лист практического испытания по дисциплине')
            par2 = doc.add_paragraph('"Оказание первой помощи"')
            # добавляем параграф
            par4 = doc.add_paragraph('Группа № ')
            # добавляем текст в третий параграф
            par4.add_run(str(Notebook.name_analis.iloc[0, 1])).bold = True  # number_grupp
            # добавляем таблицу с именами столбцов таблицы с оценками
            table5 = doc.add_table(rows=1, cols=3)
            table5.style = 'Table Grid'
            table5.cell(0, 0).text = '№ п/п'
            table5.cell(0, 1).text = 'Фамилия, имя, отчество'
            table5.cell(0, 2).text = "Оценка"
            # добавляем таблицу оценок по фамилиям
            table1 = doc.add_table(rows=Notebook.name_res.shape[0],
                                   cols=3)  # из короткой выборки имён из name_analis из
            # метода find_date()
            number = 1  # для нумерации строк таблицы
            for row in range(Notebook.name_analis.shape[0]):
                for col in range(3):
                    cell = table1.cell(row, col)
                    if col == 0:
                        cell.text = str(number)
                        number += 1
                    else:
                        cell.text = str(Notebook.name_res.iloc[row, col - 1]).title()
                        if cell.text == '5':
                            cell.text = 'отлично'
                        elif cell.text == '4':
                            cell.text = 'хорошо'
                        elif cell.text == '3':
                            cell.text = 'удовлетворительно'
                        elif cell.text == '2':
                            cell.text = 'неудовлетворительно'

            total = Notebook.name_res['оценка'].count()
            attest = Notebook.name_res['name'][Notebook.name_res['оценка'] >= 3].count()
            par5 = doc.add_paragraph('')
            # добавляем таблицу итогов по группе
            table2 = doc.add_table(rows=3, cols=2)
            table2.style = 'Table Grid'
            table2.cell(0, 0).text = 'Всего слушателей'
            table2.cell(0, 1).text = str(total)
            table2.cell(1, 0).text = 'Аттестовано'
            table2.cell(1, 1).text = str(attest)
            table2.cell(2, 0).text = 'Не аттестовано'
            table2.cell(2, 1).text = str(total - attest)
            par6 = doc.add_paragraph('')
            # добавляем таблицу с именем экзаменатора
            table3 = doc.add_table(rows=1, cols=2)
            table3.cell(0, 0).text = "Экзаменатор"
            table3.cell(0, 1).text = Notebook.name_analis.iloc[0, 5]
            par7 = doc.add_paragraph('')
            # добавляем таблицу с датой создания документа
            table4 = doc.add_table(rows=1, cols=1)
            table4.cell(0, 0).text = datetime.strftime(datetime.now(), '%d.%m.%Y')  # печать текущей даты

            try:
                doc.save(f'{Notebook.path_saving_result_and_recomendation}/results_grupp_' + str(
                    Notebook.name_analis.iloc[0, 1]) + '.docx')  # вызывается в окне выбора папки сохранения
            except:
                self.permission_denied(self)
            Notebook.grupp_analis_flag = False  # опускаем флаг анализа

    def print_result_date(self):
        """
        Создание документа "*.docx", содержащего сведения о месте экзамена, экзаменаторе
        фамилии, имена и отчества студентов, сдававших экзамен на определённую дату
        (по умолчанию, на текущую дату) с полученными оценками.
        """
        if Notebook.date_analys_flag:  # при созданном и непустом фрейме, проведённом анализе find_date()
            doc = docx.Document()
            # добавляем первый параграф (абзац)
            doc.add_paragraph(Notebook.name_analis.iloc[0, 6])  # ort из name_analis из find_date()
            # добавляем параграф
            par1 = doc.add_paragraph('Аттестационный лист практического испытания по дисциплине')
            par2 = doc.add_paragraph('"Оказание первой помощи"')
            # добавляем параграф
            par4 = doc.add_paragraph('Группа № ')
            # добавляем текст в третий параграф
            # TODO можно сделать выбор группы, если на дату сдавало больше 1 группы?
            # len(Notebook.name_analis.number_grupp.unique()) > 1
            par4.add_run(str(Notebook.name_analis.iloc[0, 1])).bold = True  # number_grupp
            # добавляем таблицу с именами столбцов таблицы с оценками
            table5 = doc.add_table(rows=1, cols=3)
            table5.style = 'Table Grid'
            table5.cell(0, 0).text = '№ п/п'
            table5.cell(0, 1).text = 'Фамилия, имя, отчество'
            table5.cell(0, 2).text = "Оценка"
            # добавляем таблицу оценок по фамилиям
            table1 = doc.add_table(rows=Notebook.name_res.shape[0],
                                   cols=3)  # из короткой выборки имён из name_analis из
            # метода find_date()
            number = 1  # для нумерации строк таблицы
            for row in range(Notebook.name_analis.shape[0]):
                for col in range(3):
                    cell = table1.cell(row, col)
                    if col == 0:
                        cell.text = str(number)
                        number += 1
                    else:
                        cell.text = str(Notebook.name_res.iloc[row, col - 1]).title()
                        if cell.text == '5':
                            cell.text = 'отлично'
                        elif cell.text == '4':
                            cell.text = 'хорошо'
                        elif cell.text == '3':
                            cell.text = 'удовлетворительно'
                        elif cell.text == '2':
                            cell.text = 'неудовлетворительно'

            total = Notebook.name_res['оценка'].count()
            attest = Notebook.name_res['name'][Notebook.name_res['оценка'] >= 3].count()
            par5 = doc.add_paragraph('')
            # добавляем таблицу итогов по группе
            table2 = doc.add_table(rows=3, cols=2)
            table2.style = 'Table Grid'
            table2.cell(0, 0).text = 'Всего слушателей'
            table2.cell(0, 1).text = str(total)
            table2.cell(1, 0).text = 'Аттестовано'
            table2.cell(1, 1).text = str(attest)
            table2.cell(2, 0).text = 'Не аттестовано'
            table2.cell(2, 1).text = str(total - attest)
            par6 = doc.add_paragraph('')
            # добавляем таблицу с именем экзаменатора
            table3 = doc.add_table(rows=1, cols=2)
            table3.cell(0, 0).text = "Экзаменатор"
            table3.cell(0, 1).text = Notebook.name_analis.iloc[0, 5]
            par7 = doc.add_paragraph('')
            # добавляем таблицу с датой создания документа
            table4 = doc.add_table(rows=1, cols=1)
            table4.cell(0, 0).text = datetime.strftime(datetime.now(), '%d.%m.%Y')  # печать текущей даты

            try:
                doc.save(f'{Notebook.path_saving_result_and_recomendation}/results_grupp_' + str(
                    Notebook.name_analis.iloc[0, 1]) + '.docx')  # вызывается в окне выбора папки сохранения
            except:
                self.permission_denied(self)
            Notebook.date_analys_flag = False  # опускаем флаг анализа

    def permission_denied(self):
        box_content = BoxLayout(orientation='vertical')
        label_warning = Label(text='Отказано в доступе\nВыберите другую папку')
        btn_yes = Button(text='Закрыть окно', on_press=(lambda *args: box_popup.dismiss()))
        box_content.add_widget(label_warning)
        box_content.add_widget(btn_yes)
        box_popup = Popup(title='Предупреждение!', content=box_content, size_hint=(0.9, 0.4))
        box_popup.open()

        #def dismiss():
        #    box_popup.dismiss()

    def false_date_popup(self):
        """всплывающее окно предупреждения о неверном формате введённой даты"""
        box_content = BoxLayout(orientation='vertical')
        label_warning = Label(text='Запишите дату в формате \n дд.мм.гггг')
        btn_yes = Button(text='Закрыть окно', on_press=(lambda *args: box_popup.dismiss()))
        box_content.add_widget(label_warning)
        box_content.add_widget(btn_yes)
        box_popup = Popup(title='Неверный формат даты!', content=box_content, size_hint=(0.9, 0.4))
        box_popup.open()

        #def dismiss():
         #   box_popup.dismiss()


class CheckExamThemes(Screen):
    """Класс предсталяет методы порождения окон всех тем экзамена;
    содержит флаги выбора отдельных тем экзамена, накапливает оценки и данные для отображения имени и
    группы на окнах экзаменационных тем; методы всплывающих окон предупреждения о невывполненных
    заданиях; инициализирует расчёт предварительной и итоговой ощенок"""

    Builder.load_file('CheckExamThemes.kv')

    total_bewertung = ''  # результат записывает результаты всех этапов экзамена
    first_visio_flag = True  # on_switch == False
    recreation_flag = True
    arterial_flag = True
    venosus_flag = True
    cpr_flag = True
    frakt_zakr_flag = True
    frakt_otkr_flag = True
    heiml_otro_flag = True
    heiml_down_flag = True
    transport_flag = True

    name_student = ''  # должно отражаться в каждом скрине
    number_grupp = 0  # номер группы

    result_flag = 0  # хранение кол-ва выполненных заданий

    def build_first_visio(self):
        """Порождение и переход в окно экзамена """
        if self.first_visio_flag:
            FirstVisio.name_student = Notebook.name_student  # устанавливаем имя студента
            FirstVisio.summ_first_visio = Notebook.summ_result  # передаём баллы
            FirstVisio.number_grupp = Notebook.number_grupp  # передаём номер группы
            MainApp.sm.add_widget(FirstVisio(name='first_visio'))
            MainApp.sm.current = 'first_visio'
            self.result_flag += 1

    def build_recreation(self):
        """Порождение и переход в окно экзамена """
        if self.recreation_flag:
            RecreationState.name_student = Notebook.name_student
            RecreationState.recreate_summ = Notebook.summ_result
            RecreationState.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(RecreationState(name='recreate_state'))
            MainApp.sm.current = 'recreate_state'
            self.result_flag += 1

    def build_arterial(self):
        """Порождение и переход в окно экзамена """
        if self.arterial_flag:
            Arterial.name_student = Notebook.name_student
            Arterial.arterial_summ = Notebook.summ_result
            Arterial.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(Arterial(name='arterial'))
            MainApp.sm.current = 'arterial'
            self.result_flag += 1

    def build_venosus(self):
        """Порождение и переход в окно экзамена """
        if self.venosus_flag:
            Venosus.name_student = Notebook.name_student
            Venosus.venosus_summ = Notebook.summ_result
            Venosus.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(Venosus(name='venosus'))
            MainApp.sm.current = 'venosus'
            self.result_flag += 1

    def build_fraktura_zakr(self):
        """Порождение и переход в окно экзамена """
        if self.frakt_zakr_flag:
            FracturaZakr.name_student = Notebook.name_student
            FracturaZakr.fractura_zakr_summ = Notebook.summ_result
            FracturaZakr.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(FracturaZakr(name='fractura_zakr'))
            MainApp.sm.current = 'fractura_zakr'
            self.result_flag += 1

    def build_fraktura_otkr(self):
        """Порождение и переход в окно экзамена """
        if self.frakt_otkr_flag:
            FracturaOtkr.name_student = Notebook.name_student
            FracturaOtkr.fraktura_otkr_summ = Notebook.summ_result
            FracturaOtkr.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(FracturaOtkr(name='fractura_otkr'))
            MainApp.sm.current = 'fractura_otkr'
            self.result_flag += 1

    def build_cpr(self):
        """Порождение и переход в окно экзамена """
        if self.cpr_flag:
            CPR.name_student = Notebook.name_student
            CPR.cpr_summ = Notebook.summ_result
            CPR.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(CPR(name='cpr'))
            MainApp.sm.current = 'cpr'
            self.result_flag += 1

    def build_heiml_orto(self):
        """Порождение и переход в окно экзамена """
        if self.heiml_otro_flag:
            HeimlichVertical.name_student = Notebook.name_student
            HeimlichVertical.hieml_vert_summ = Notebook.summ_result
            HeimlichVertical.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(HeimlichVertical(name='heimlich_vertical'))
            MainApp.sm.current = 'heimlich_vertical'
            self.result_flag += 1

    def build_heiml_down(self):
        """Порождение и переход в окно экзамена """
        if self.heiml_down_flag:
            HeimlichHorizontal.name_student = Notebook.name_student
            HeimlichHorizontal.heiml_hor_summ = Notebook.summ_result
            HeimlichHorizontal.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(HeimlichHorizontal(name='heimlich_horizontal'))
            MainApp.sm.current = 'heimlich_horizontal'
            self.result_flag += 1

    def build_transport(self):
        """Порождение и переход в окно экзамена """
        if self.transport_flag:
            Transport.name_student = Notebook.name_student
            Transport.transport_summ = Notebook.summ_result
            Transport.number_grupp = Notebook.number_grupp
            MainApp.sm.add_widget(Transport(name='transport'))
            MainApp.sm.current = 'transport'
            self.result_flag += 1

    def build_dopolnitelno(self):
        """Порождение и переход в окно дополнительных вопросов"""
        Dopolnitelno.name_student = Notebook.name_student
        Dopolnitelno.number_grupp = Notebook.number_grupp
        MainApp.sm.add_widget(Dopolnitelno(name='dopolnitelno'))
        MainApp.sm.current = 'dopolnitelno'

    def close_widget(self):
        self.result_flag = 0
        MainApp.sm.remove_widget(self)

    def bewertung_popup(self):
        """Всплывающее окно предупреждения о незаконченных темах;
        расчёт предварительной оценки"""
        if self.result_flag == 10:
            Notebook.bewertung_result(self=Notebook)  # расчёт оенки
            self.choose_result_builder()
        else:
            self.warning_popup()

    def warning_popup(self):
        """Всплывающее окно незаконченных заданий"""
        box_content = BoxLayout(orientation='vertical')
        label_warning = Label(text='Вернитесь к выполнению заданий!'.center(20))
        box_intra_content = BoxLayout(orientation='horizontal')
        btn_yes = Button(text='Вернуться', on_press=(lambda *args: box_popup.dismiss()))
        btn_close = Button(text='Прекратить \nэкзамен', on_press=lambda *args: no_exam())
        box_intra_content.add_widget(btn_yes)
        box_intra_content.add_widget(btn_close)
        box_content.add_widget(label_warning)
        box_content.add_widget(box_intra_content)
        box_popup = Popup(title='Предупреждение!', content=box_content, size_hint=(0.8, 0.4))
        box_popup.open()

        #def dismiss():
         #   box_popup.dismiss()

        def no_exam():
            """экзамен прекращается без записи результатов  журнал"""
            box_popup.dismiss()
            self.close_widget()

    def choose_result_builder(self):
        """всплывающее окно результата и выбор дополнительных вопросов"""
        box_content = BoxLayout(orientation='vertical')
        box_intra_content = BoxLayout(orientation='horizontal')
        label_name = Label(text=self.name_student)
        label_grupp = Label(text='группа № ' + str(self.number_grupp))
        label_result = Label(text='Штрафных баллов: ' + str(Notebook.summ_result))
        label_bewertung = Label(text='Предварительная оценка: ' + str(Notebook.bewertung))
        btn_yes = Button(text='ДопВопросы',
                         on_press=(lambda *args: to_dopolnitelno()))
        #btn_yes.text_size = btn_yes.size
        btn_no = Button(text='Закончить экзамен', on_press=(lambda *args: no_dopolnitelno()))
        box_content.add_widget(label_name)
        box_content.add_widget(label_grupp)
        box_content.add_widget(label_result)
        box_content.add_widget(label_bewertung)
        box_intra_content.add_widget(btn_yes)
        box_intra_content.add_widget(btn_no)
        box_content.add_widget(box_intra_content)
        box_popup = Popup(title='Предварительные результаты',
                          content=box_content,
                          size_hint=(0.8, 0.4))
        box_popup.open()

        def to_dopolnitelno():
            """экзамен данного студента окончен, переход к допВопросам"""
            self.disabled_btn()
            box_popup.dismiss()

        def no_dopolnitelno():
            """экзамен данного студента окончен, нет перехода к допВопросам"""
            Notebook.data['summ_fehler'] = Notebook.summ_result  # запись в словарь
            Notebook.csv_writer(self=Notebook)  # записываем в журнал результат
            box_popup.dismiss()
            self.close_widget()

    def disabled_btn(self):  # блокировка кнопки "Дополнительные вопросы"
        self.ids.btn_dopolnitelno.disabled = False


class FirstVisio(Screen):
    Builder.load_file('FirstVisio.kv')

    name_student = ''
    number_grupp = 0
    lbl_name_student = ObjectProperty()
    end_ohne_fehler = ObjectProperty()
    summ_first_visio = 0  # результат этапа
    color_btn_choose = True
    lst_fehler_1 = ['Тактильное раздражение', 'Вызов помощи', 'Обращение голосом']
    lst_fehler_2 = ['Запрокидывание головы', 'Общий осмотр']

    def display_result(self):
        self.display.text = str(self.summ_first_visio)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку завершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.summ_first_visio += 1  # увеличение результата этапа
            self.display_result()  # отображение езультата этапа на дисплее штрафов
            Notebook.data[f"Первичный осмотр. {instance.text}"] = 1  # запись в словарь значения ошибки
        elif instance.text in self.lst_fehler_2:
            self.summ_first_visio += 2  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Первичный осмотр. {instance.text}"] = 2
        else:
            self.summ_first_visio += 3  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Первичный осмотр. {instance.text}"] = 3

    def close_widget(self):
        Notebook.summ_result = self.summ_first_visio  # передаём результат этапа в notebook
        MainApp.sm.remove_widget(self)


class RecreationState(Screen):
    Builder.load_file('RecreationState.kv')

    recreate_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0
    lst_fehler_1 = ['Выпрямление руки', 'Сгибание руки', 'Сгибание ноги', 'Поворот тела']

    def display_result(self):
        self.display.text = str(self.recreate_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.recreate_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Восстановительное положение. {instance.text}"] = 1
        else:
            self.recreate_summ += 3  # 3 критическая == сумма их должна быть больше суммы некритических
            self.display_result()
            Notebook.data[f"Восстановительное положение. {instance.text}"] = 3

    def close_widget(self):
        Notebook.summ_result = self.recreate_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class Arterial(Screen):
    Builder.load_file('Arterial.kv')

    arterial_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['На тканевую подложку', 'Записка', 'Иммобилизация', 'Перчатки']

    def display_result(self):
        self.display.text = str(self.arterial_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.arterial_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Артериальное кровотечение. {instance.text}"] = 1
        else:
            self.arterial_summ += 3
            self.display_result()
            Notebook.data[f"Артериальное кровотечение. {instance.text}"] = 3

    def close_widget(self):
        Notebook.summ_result = self.arterial_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class Venosus(Screen):
    Builder.load_file('Venosus.kv')

    venosus_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['Бинт раскатывается', 'Холод']

    def display_result(self):
        self.display.text = str(self.venosus_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.venosus_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Венозное кровотечение. {instance.text}"] = 1
        else:
            self.venosus_summ += 2
            self.display_result()
            Notebook.data[f"Венозное кровотечение. {instance.text}"] = 2

    def close_widget(self):
        Notebook.summ_result = self.venosus_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class FracturaZakr(Screen):
    Builder.load_file('FrakturaZakr.kv')

    fractura_zakr_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['Вызов помощника', 'Холод']

    lst_fehler_3 = ['Изгибание шины Крамера', 'Захват суставов']

    def display_result(self):
        self.display.text = str(self.fractura_zakr_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.fractura_zakr_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Закрытый перелом. {instance.text}"] = 1
        elif instance.text in self.lst_fehler_3:
            self.fractura_zakr_summ += 3
            self.display_result()
            Notebook.data[f"Закрытый перелом. {instance.text}"] = 3
        else:
            self.fractura_zakr_summ += 2
            self.display_result()
            Notebook.data[f"Закрытый перелом. {instance.text}"] = 2

    def close_widget(self):
        Notebook.summ_result = self.fractura_zakr_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class FracturaOtkr(Screen):
    Builder.load_file('FrakturaOtkr.kv')

    fraktura_otkr_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['Холод']

    lst_fehler_3 = ['Повязка на рану', 'Захват суставов', 'Изгибание шины Крамера', 'Длина фиксации шины']

    def display_result(self):
        self.display.text = str(self.fraktura_otkr_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.fraktura_otkr_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Открытый перелом. {instance.text}"] = 1
        elif instance.text in self.lst_fehler_3:
            self.fraktura_otkr_summ += 3
            self.display_result()
            Notebook.data[f"Открытый перелом. {instance.text}"] = 3
        else:
            self.fraktura_otkr_summ += 2
            self.display_result()
            Notebook.data[f"Открытый перелом. {instance.text}"] = 2

    def close_widget(self):
        Notebook.summ_result = self.fraktura_otkr_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class CPR(Screen):
    Builder.load_file('CPR.kv')

    cpr_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_6 = ['Переломы']
    lst_fehler_3 = ['Вызов помощи', 'Контроль дыхания', 'Руки прямые / кисти в замок',
                    'Частота / глубина / количество']

    def display_result(self):
        self.display.text = str(self.cpr_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_3:
            self.cpr_summ += 3  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Сердечно-лёгочная реанимация. {instance.text}"] = 3
        elif instance.text in self.lst_fehler_6:
            self.cpr_summ += 6
            self.display_result()
            Notebook.data[f"Сердечно-лёгочная реанимация. {instance.text}"] = 6
        else:
            self.cpr_summ += 2
            self.display_result()
            Notebook.data[f"Сердечно-лёгочная реанимация. {instance.text}"] = 2

    def close_widget(self):
        Notebook.summ_result = self.cpr_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class Transport(Screen):
    Builder.load_file('Transport.kv')

    transport_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['Организовать помощь', 'Таз удерживается']

    def display_result(self):
        self.display.text = str(self.transport_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.transport_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Вынос и транспортировка. {instance.text}"] = 1
        elif instance.text == 'Голова и шея фиксированы':
            self.transport_summ += 2
            self.display_result()
            Notebook.data[f"Вынос и транспортировка. {instance.text}"] = 2
        else:
            self.transport_summ += 3
            self.display_result()
            Notebook.data[f"Вынос и транспортировка. {instance.text}"] = 3

    def close_widget(self):
        Notebook.summ_result = self.transport_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class Dopolnitelno(Screen):
    Builder.load_file('Dopolnitelno.kv')
    dopolnitel_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    def result(self, instance):
        self.dopolnitel_summ += 2  # уменьшение результата этапа
        self.display.text = str(self.dopolnitel_summ)  # отражаем на дисплее штрафов
        Notebook.data[f"Корректирующие баллы. {instance.text}"] = 2

    def destroy_windows(self):  # всплывающее окно с результатом, уничтожение всех окон экзамена, возврат в меню
        Notebook.summ_result -= self.dopolnitel_summ  # передаём результат этапа в ноутбук
        Notebook.data['summ_fehler'] = Notebook.summ_result  # запись в словарь результатов
        Notebook.bewertung_result(self=Notebook)  # расчёт оенки
        Notebook.csv_writer(self=Notebook)  # запись в журнал
        MainApp.sm.remove_widget(self)

    def result_popup(self):
        popup_content = BoxLayout(orientation='vertical')
        label_result = Label(text='Окончательный результат: ' + str(Notebook.bewertung))
        btn_ok = Button(text='Закрыть', on_press=(lambda *args: box_popup.dismiss()), size_hint=(1, 0.4))
        popup_content.add_widget(label_result)
        popup_content.add_widget(btn_ok)
        box_popup = Popup(title='Итог экзамена',
                          content=popup_content,
                          size_hint=(0.8, 0.4))
        box_popup.open()


class HeimlichVertical(Screen):
    Builder.load_file('HeimlichVertical.kv')
    hieml_vert_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_1 = ['Предупреждение', 'Эффективность до 5-й попытки']
    lst_fehler_2 = ['Наклон туловища', 'Плечи ниже', 'Кисти / место / локти']
    lst_fehler_3 = ['Таз обозначен', 'Спина прижата']

    def display_result(self):
        self.display.text = str(self.hieml_vert_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_1:
            self.hieml_vert_summ += 1  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха стоя. {instance.text}"] = 1
        elif instance.text in self.lst_fehler_2:
            self.hieml_vert_summ += 2
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха стоя. {instance.text}"] = 2
        else:
            self.hieml_vert_summ += 3
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха стоя. {instance.text}"] = 3

    def close_widget(self):
        Notebook.summ_result = self.hieml_vert_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class HeimlichHorizontal(Screen):
    Builder.load_file('HeimlichHorizontal.kv')

    heiml_hor_summ = 0  # результат этапа
    name_student = ''
    number_grupp = 0

    lst_fehler_2 = ['Направляющая рука установлена', 'Ударная рука скользит']
    lst_fehler_3 = ['Таз обозначен', 'Голова повёрнута']

    def display_result(self):
        self.display.text = str(self.heiml_hor_summ)  # отражаем на дисплее штрафов
        self.ids.end_ohne_fehler.disabled = True  # блокировать кнопку заввершения

    def result(self, instance):
        if instance.text in self.lst_fehler_2:
            self.heiml_hor_summ += 2  # увеличение результата этапа
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха лежащему. {instance.text}"] = 2
        elif instance.text in self.lst_fehler_3:
            self.heiml_hor_summ += 3
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха лежащему. {instance.text}"] = 3
        else:
            self.heiml_hor_summ += 1
            self.display_result()
            Notebook.data[f"Манёвр Геймлиха лежащему. {instance.text}"] = 1

    def close_widget(self):
        Notebook.summ_result = self.heiml_hor_summ  # передаём результат этапа в чек
        MainApp.sm.remove_widget(self)


class MySwiper(Screen):
    Builder.load_file('myswiper.kv')

    images_dictionary = {"Начальный экран": "instruction_images/1.png",
                         'Справка': "instruction_images/spravka.png",
                         'Главный экран': "instruction_images/главный экран.png",
                         "Место проведения": "instruction_images/выбор места.png",
                         "Предупреждение имени": "instruction_images/предупреждение имени студента.png",
                         "Выбор частичного экзамена": "instruction_images/к выбору тем.png",
                         "Выбор темы": "instruction_images/выбор темы.png",
                         "Темы экзамена": "instruction_images/меню тем экзамена.png",
                         "Тема": "instruction_images/тема.png",
                         "Тема с ошибками": "instruction_images/тема с ошибкой.png",
                         "Темы отмечены": "instruction_images/темы отмечены.png",
                         "Незавершённый экзамен": "instruction_images/предупреждение незаврешённого экзамена.png",
                         "Все темы пройдены": "instruction_images/все темы пройдены.png",
                         "Предварительная оценка": "instruction_images/предварительная оценка.png",
                         "Активны дополнительные темы": "instruction_images/активны допВопросы.png",
                         "Дополнительные темы": "instruction_images/допВопросы.png",
                         "Окончательная оценка": "instruction_images/окончательная оценка.png",
                         "К следующему студенту": "instruction_images/к следующему студенту.png",
                         "Итоговый документ": "instruction_images/документ итогов.png",
                         "Экран итогов": "instruction_images/итоги.png",
                         "Итог группы": "instruction_images/итоги группы.png",
                         "Поиск по имени": "instruction_images/итоги поиска по имени.png",
                         "Выбор папки сохранения": "instruction_images/выбор папки.png"}

    def spinner_my_swiper_clicked(self):
        self.ids.image_swiper.source = self.images_dictionary[self.ids.spinner_swiper_id.text]

    def exit_swipe(self):
        MainApp.sm.remove_widget(self)
        MainApp.sm.current = 'spinner'


class MyLayout(Screen):
    path = 'spravka_summary.txt'
    Builder.load_file('spin.kv')

    def close_window(self):
        MainApp.sm.remove_widget(self)

    def spinner_clicked(self):
        if self.ids.spinner_id.text == 'Справка о работе программы':
            with open(self.path, 'r') as file:
                result = file.read()
                self.ids.down_text_scroll.text = result
        elif self.ids.spinner_id.text == 'Система оценок':
            with open('about.txt', 'r') as file:
                result = file.read()
                self.ids.down_text_scroll.text = result

    def build_my_swiper(self):
        MainApp.sm.add_widget(MySwiper(name='myswiper'))
        MainApp.sm.current = 'myswiper'


class indoManager(ScreenManager):
    pass

kv = Builder.load_file('main.kv')

class MainApp(App):
    sm = indoManager()

    def build(self):
        #self.sm.add_widget(StartScreen(name='start_screen'))
        #return self.sm
        return kv


class Notebook:
    res_name = pd.DataFrame()  # фрейм результато поиска по имени
    result = pd.DataFrame()  # главный фрейм
    grupp_analis_flag = False  # флаг уже проведённого анализа данных по группе
    name_fehler_flag = False  # флаг уже проведённого анализа данных по имени
    date_analys_flag = False  # флаг уже проведённого анализа данных по дате
    path_saving_result_and_recomendation = ''
    res_date = pd.DataFrame()
    name_analis = pd.DataFrame()  # требуется для оформления и вывода файла результатов
    name_res = pd.DataFrame()  # требуется для оформления и вывода файла результатов
    res_grupp = pd.DataFrame()  # требуется для оформления и вывода файла результатов
    part_MAXIMUM = 127  # максимальная сумма баллов при частичном экзамене
    name_student = ''  # имя студента
    number_grupp = 0  # номер группы
    prufer = ['Трус Алексей Эдуардович']
    member_prufer = ''
    ort = ['ЧОУ ДПО Учебный центр ПАО Газпром']
    member_ort = ''
    summ_result = 0  # общая сумма баллов
    MAXIMUM = 127  # максимальное кол-во баллов при полном экзамене TODO проверь снова сколько баллов
    bewertung = 0  # итоговая оценка

    choose_themes_flag = bool  # частичный экзамен == True

    z = datetime.now()
    date_prufung = datetime.strftime(z, '%d.%m.%Y')

    id_student = 0
    # TODO вставить ключ и булево поле частичного экзамена для анализа после тестирования анализа?
    data = {'ID': id_student, 'name': name_student, 'number_grupp': number_grupp, 'date': date_prufung,
            'оценка': bewertung, 'summ_fehler': 0, 'prufer': member_prufer, 'ort': member_ort,
            'Первичный осмотр. Безопасное приближение': 0, "Первичный осмотр. Обращение голосом": 0,
            "Первичный осмотр. Тактильное раздражение": 0, "Первичный осмотр. Вызов помощи": 0,
            "Первичный осмотр. Запрокидывание головы": 0,
            'Первичный осмотр. Выдвижение нижней челюсти': 0, 'Первичный осмотр. Определение дыхания': 0,
            'Первичный осмотр. Общий осмотр': 0, 'Первичный осмотр. Контроль дыхания непрерывно': 0,
            'Восстановительное положение. Выпрямление руки': 0, 'Восстановительное положение. Сгибание руки': 0,
            'Восстановительное положение. Сгибание ноги': 0, 'Восстановительное положение. Поворот тела': 0,
            'Восстановительное положение. Контроль дыхания после осмотра': 0,
            'Артериальное кровотечение. Прижатие на протяжении': 0, 'Артериальное кровотечение. Место наложения': 0,
            'Артериальное кровотечение. На тканевую подложку': 0, 'Артериальное кровотечение. Записка': 0,
            'Артериальное кровотечение. Иммобилизация': 0, 'Артериальное кровотечение. Перчатки': 0,
            'Венозное кровотечение. Перчатки': 0, 'Венозное кровотечение. Бинт раскатывается': 0,
            'Венозное кровотечение. Давящий валик': 0, 'Венозное кровотечение. Холод': 0,
            'Сердечно-лёгочная реанимация. Вызов помощи': 0, 'Сердечно-лёгочная реанимация. Контроль дыхания': 0,
            'Сердечно-лёгочная реанимация. Одежда / ремень': 0,
            'Сердечно-лёгочная реанимация. Руки прямые / кисти в замок': 0,
            'Сердечно-лёгочная реанимация. Переломы': 0,
            'Сердечно-лёгочная реанимация. Частота / глубина / количество': 0,
            'Сердечно-лёгочная реанимация. Голова запрокинута / нос закрыт': 0,
            'Сердечно-лёгочная реанимация. Челюсть / обзор грудной': 0,
            'Сердечно-лёгочная реанимация. Скорость вдоха / количество': 0,
            'Сердечно-лёгочная реанимация. Интервалы / длительность': 0,
            'Закрытый перелом. Вызов помощника': 0, 'Закрытый перелом. Изгибание шины Крамера': 0,
            'Закрытый перелом. Длина фиксации шины': 0,
            'Закрытый перелом. Захват суставов': 0, 'Закрытый перелом. Холод': 0,
            'Открытый перелом. Повязка на рану': 0, 'Открытый перелом. Перчатки': 0,
            'Открытый перелом. Изгибание шины Крамера': 0,
            'Открытый перелом. Длина фиксации шины': 0, 'Открытый перелом. Захват суставов': 0,
            'Открытый перелом. Холод': 0, 'Манёвр Геймлиха стоя. Предупреждение': 0,
            'Манёвр Геймлиха стоя. Наклон туловища': 0,
            'Манёвр Геймлиха стоя. Таз обозначен': 0, 'Манёвр Геймлиха стоя. Плечи ниже': 0,
            'Манёвр Геймлиха стоя. Спина прижата': 0, 'Манёвр Геймлиха стоя. Кисти / место / локти': 0,
            'Манёвр Геймлиха стоя. Эффективность до 5-й попытки': 0,
            'Манёвр Геймлиха лежащему. Таз обозначен': 0, 'Манёвр Геймлиха лежащему. Направляющая рука установлена': 0,
            'Манёвр Геймлиха лежащему. Ударная рука скользит': 0, 'Манёвр Геймлиха лежащему. Голова повёрнута': 0,
            'Манёвр Геймлиха лежащему. Эффективность до 5-й попытки': 0,
            'Вынос и транспортировка. Организовать помощь': 0, 'Вынос и транспортировка. Голова и шея фиксированы': 0,
            'Вынос и транспортировка. Таз удерживается': 0, 'Вынос и транспортировка. Контроль головы при выносе': 0,
            'Корректирующие баллы. Шейный фиксирующий воротник': 0,
            'Корректирующие баллы. Дыхательные: плёнка / маска / мешок': 0,
            'Корректирующие баллы. Окклюзионная повязка': 0, 'Корректирующие баллы. Автоматический дефибриллятор': 0,
            'Корректирующие баллы. Индикатор сердечного ритма': 0,
            'Корректирующие баллы. Транспортировка на носилках по лестницам': 0}  # словарь итоговых данных

    def bewertung_result(self):  # подсчёт итогового результата. вызывается дважды - для предвариьтельного и
        # окончательного подсчётов
        # TODO записать кол-во ошибок в словарь. нет переменной! не надо!? не отсюда!
        AdministerHuman.choose_themes_widget(self=AdministerHuman)  # флаги каждой темы поднять
        if self.choose_themes_flag:  # если выбран частичный экзамен
            tempo = self.part_MAXIMUM - self.summ_result
            if (self.part_MAXIMUM * 0.85) <= tempo:
                self.bewertung = 5
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif (self.part_MAXIMUM * 0.7) <= tempo <= (self.part_MAXIMUM * 0.84):
                self.bewertung = 4
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif (self.part_MAXIMUM * 0.54) <= tempo <= (self.part_MAXIMUM * 0.69):
                self.bewertung = 3
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif tempo < self.part_MAXIMUM * 0.54:
                self.bewertung = 2
                self.data['оценка'] = self.bewertung
                return self.bewertung
        else:  # если выбран полный экзамен
            tempo = self.MAXIMUM - self.summ_result
            if (self.MAXIMUM * 0.85) <= tempo:
                self.bewertung = 5
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif (self.MAXIMUM * 0.7) <= tempo <= (self.MAXIMUM * 0.84):
                self.bewertung = 4
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif (self.MAXIMUM * 0.54) <= tempo <= (self.MAXIMUM * 0.69):
                self.bewertung = 3
                self.data['оценка'] = self.bewertung
                return self.bewertung
            elif tempo < self.MAXIMUM * 0.54:
                self.bewertung = 2
                self.data['оценка'] = self.bewertung
                return self.bewertung

    def config_reader(self):
        with open('config.txt', 'r') as conf_file:
            res_read = conf_file.read().split(',')
            readable_id = int(res_read[0])
            self.member_prufer = res_read[1]
            self.member_ort = res_read[2]
            self.id_student = readable_id
            return self.id_student, self.member_ort, self.member_prufer

    def config_writer(self):
        with open('config.txt', 'w') as conf_file:
            str_for_config = ','.join([str(self.id_student), self.prufer[-1], self.ort[-1]])
            conf_file.write(str_for_config)
        self.data['ID'] = self.id_student

    def id_student_plus(self):
        self.id_student += 1
        return self.id_student

    def csv_writer(self):
        """ Формирование итогового файла хранения результатов. """
        if os.path.exists('students_data.csv'):
            with open('students_data.csv', 'a') as scv_file:
                columns = self.data.keys()
                writer = csv.DictWriter(scv_file, fieldnames=columns)
                writer.writerow(self.data)
        else:
            with open('students_data.csv', 'w') as scv_file:
                columns = self.data.keys()
                writer = csv.DictWriter(scv_file, fieldnames=columns)
                writer.writeheader()
                writer.writerow(self.data)


if __name__ == '__main__':
    MainApp().run()
